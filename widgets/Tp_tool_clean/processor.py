# processor.py (backend logic)

import tempfile
import requests
from docx import Document
import openpyxl
import pdfplumber
import pandas as pd

# ——— Azure OpenAI config ———
# Expect these to be set by the Streamlit frontend via secrets or environment variables
API_KEY = None  # to be set by frontend
API_ENDPOINT = None  # to be set by frontend


def configure(api_key: str, api_endpoint: str):
    global API_KEY, API_ENDPOINT
    API_KEY = api_key
    API_ENDPOINT = api_endpoint


def get_llm_response_azure(prompt: str, context: str) -> str:
    headers = {"Content-Type": "application/json", "api-key": API_KEY}
    system_msg = (
        "You are an expert on Transfer Pricing and financial analysis. "
        "Use the information in the following context to answer the user's question. "
        "Assign the greatest priority to the information that you gather from the financial analysis and the interview transcript. "
        "If asked something not covered in this data, you may search the web."
        "Ensure your analysis is CONCISE, SHARP, in paragraph form, and not long. Never use bullet points. "
        "DO NOT INCLUDE MARKDOWN FORMATTING OR # SIGNS. Keep it to 200-300 words, maintain a professional tone. "
        "Make sure to include direct sources and citations for the data you use for your decisions. Also include your reasoning for conclusions in brackets ()."
        "If something is from the transcript or financial statement, include that citation in brackets with a URL to the specific section. Likewise include a URL to the relevant website if the information you got was from searching the internet. "
        "You **may** consider the OECD guidelines below as helpful targets, but do NOT structure your response around them.\n\n"
    )
    messages = [
        {"role": "system", "content": system_msg},
        {"role": "user", "content":  context + prompt}
    ]
    resp = requests.post(API_ENDPOINT, headers=headers, json={"messages": messages})
    resp.raise_for_status()
    return resp.json()["choices"][0]["message"]["content"].strip()


def load_transcript(file) -> str:
    doc = Document(file)
    return "\n".join(p.text for p in doc.paragraphs)


def load_pdf(file) -> str:
    pages, tables = [], []
    with pdfplumber.open(file) as pdf:
        for i, page in enumerate(pdf.pages, start=1):
            text = page.extract_text() or ""
            pages.append(f"--- Page {i} ---\n{text}")
            for table in page.extract_tables():
                df = pd.DataFrame(table[1:], columns=table[0])
                tables.append(f"--- Page {i} table ---\n" + df.to_csv(index=False))
    return "\n\n".join(pages + tables)


def load_guidelines(file) -> str:
    return file.read().decode("utf-8").strip()


def load_and_annotate_replacements(excel_file, context: str) -> dict:
    wb = openpyxl.load_workbook(excel_file)
    ws = wb.active
    replacements = {}
    for row in ws.iter_rows(min_row=2, max_col=6):
        _, cell_c, cell_d, cell_e, _ = row[1], row[2], row[3], row[4], row[5]
        placeholder = cell_d.value
        if not placeholder:
            continue
        ph = str(placeholder)
        if cell_e.value and str(cell_e.value).strip():
            raw = str(cell_e.value)
            for k, v in replacements.items():
                raw = raw.replace(k, v)
            value = get_llm_response_azure(raw, context)
        else:
            value = str(cell_c.value or "")
        ws.cell(row=cell_d.row, column=7, value=value)
        replacements[ph] = value
    wb.save(excel_file)
    return replacements


def collapse_runs(paragraph):
    from docx.oxml.ns import qn
    text = "".join(r.text for r in paragraph.runs)
    for r in reversed(paragraph.runs):
        r._element.getparent().remove(r._element)
    paragraph.add_run(text)


def replace_in_paragraph(p, replacements):
    collapse_runs(p)
    for run in p.runs:
        for ph, val in replacements.items():
            if ph in run.text:
                run.text = run.text.replace(ph, val)


def replace_placeholders(doc: Document, replacements: dict):
    from docx.oxml.ns import qn
    seen = False
    br_tag = qn('w:br')
    for p in doc.paragraphs:
        if not seen:
            for r in p.runs:
                for br in r._element.findall(br_tag):
                    if br.get(qn('w:type')) == 'page':
                        seen = True
                        break
                if seen: break
            if not seen: continue
        replace_in_paragraph(p, replacements)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_in_paragraph(p, replacements)
    for sec in doc.sections:
        if sec.header:
            for p in sec.header.paragraphs:
                replace_in_paragraph(p, replacements)
        if sec.footer:
            for p in sec.footer.paragraphs:
                replace_in_paragraph(p, replacements)


def replace_first_page_placeholders(doc: Document, replacements: dict):
    from docx.oxml.ns import qn
    seen = False
    br_tag = qn("w:br"); typ = qn("w:type")
    for p in doc.paragraphs:
        replace_in_paragraph(p, replacements)
        for r in p.runs:
            for child in r._element:
                if child.tag == br_tag and child.get(typ) == "page":
                    seen = True
                    break
            if seen: break
        if seen: break


def process_and_fill(files: dict) -> str:
    # files: {'guidelines','transcript','pdf','excel','template'}
    ctx = load_guidelines(files['guidelines'])
    ctx += "\n\n" + load_transcript(files['transcript'])
    ctx += "\n\n" + load_pdf(files['pdf'])
    replacements = load_and_annotate_replacements(files['excel'], ctx)
    doc = Document(files['template'])
    replace_first_page_placeholders(doc, replacements)
    replace_placeholders(doc, replacements)
    out = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(out.name)
    return out.name
